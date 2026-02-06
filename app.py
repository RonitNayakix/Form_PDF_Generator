import streamlit as st
import streamlit.components.v1 as components

import os
import tempfile
import re

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas



# ======================================================
# CONFIG
# ======================================================
st.set_page_config(page_title="Dynamic PDF Generator", layout="centered")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
ACTIVE_TEMPLATE = os.path.join(TEMPLATE_DIR, "active_template.docx")

os.makedirs(TEMPLATE_DIR, exist_ok=True)

ADMIN_PASSWORD = st.secrets.get("ADMIN_PASSWORD", "admin123")

# ======================================================
# UTILS
# ======================================================
def extract_placeholders_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return sorted(set(re.findall(r"{{(.*?)}}", text)))


def docx_to_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def replace_placeholders(text, data):
    for k, v in data.items():
        text = text.replace(f"{{{{{k}}}}}", str(v))
    return text


def generate_pdf(content, output_path):
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    y = height - 50

    for line in content.split("\n"):
        c.drawString(40, y, line)
        y -= 18
        if y < 40:
            c.showPage()
            y = height - 50

    c.save()
    
def extract_docx_full_text(docx_path):
    doc = Document(docx_path)
    content = []

    for p in doc.paragraphs:
        content.append(p.text)

    for table in doc.tables:
        content.append("")
        for row in table.rows:
            content.append(" | ".join(cell.text for cell in row.cells))
        content.append("")

    return "\n".join(content)

def render_docx_preview(docx_path):
    raw_text = extract_docx_full_text(docx_path)

    # Highlight placeholders {{Field}}
    highlighted = re.sub(
        r"\{\{(.*?)\}\}",
        r"<strong style='color:black;'>{{\1}}</strong>",
        raw_text
    )

    # Convert newlines to <br>
    highlighted = highlighted.replace("\n", "<br>")

    st.markdown(
        f"""
        <div style="
            background-color:white;
            color:black;
            padding:16px;
            border-radius:8px;
            border:1px solid #ddd;
            font-family:Arial, sans-serif;
            line-height:1.6;
        ">
            {highlighted}
        </div>
        """,
        unsafe_allow_html=True
    )


# ======================================================
# UI
# ======================================================
st.title("📄 Dynamic PDF Generator")

tabs = st.tabs(["👤 User", "🔐 Admin"])

# ======================================================
# ADMIN TAB
# ======================================================
with tabs[1]:
    st.subheader("🔐 Admin Panel")

    password = st.text_input("Admin Password", type="password")

    if password == ADMIN_PASSWORD:

        st.markdown("### 📂 Template Overview")

        if os.path.exists(ACTIVE_TEMPLATE):
            fields = extract_placeholders_from_docx(ACTIVE_TEMPLATE)

            st.success("✅ Active template found")
            st.write("**File:** active_template.docx")
            st.write("**Extracted Fields:**")
            st.code(", ".join(fields))
            st.markdown("### 🖼 Template Preview")
            render_docx_preview(ACTIVE_TEMPLATE)

            with open(ACTIVE_TEMPLATE, "rb") as f:
                st.download_button(
                    "⬇️ Download Current Template",
                    f,
                    file_name="active_template.docx"
                )
        else:
            st.warning("❌ No template uploaded yet")

        st.divider()

        st.markdown("### ⬆️ Upload / Replace Template")

        uploaded_docx = st.file_uploader(
            "Upload DOCX ({{FieldName}} format)",
            type=["docx"]
        )

        if uploaded_docx:
            with open(ACTIVE_TEMPLATE, "wb") as f:
                f.write(uploaded_docx.read())

            st.success("✅ Template saved permanently")
            st.rerun()

    else:
        st.warning("Admin access only")

# ======================================================
# USER TAB
# ======================================================
with tabs[0]:
    st.subheader("👤 Fill Details")

    if not os.path.exists(ACTIVE_TEMPLATE):
        st.warning("Template not available. Please contact admin.")
    else:
        fields = extract_placeholders_from_docx(ACTIVE_TEMPLATE)

        pdf_name = st.text_input("📄 PDF File Name (without .pdf)")

        user_data = {}

        with st.form("user_form"):
            for field in fields:
                user_data[field] = st.text_input(field)

            submitted = st.form_submit_button("🚀 Generate PDF")

        if submitted:
            if not pdf_name.strip():
                st.error("❌ PDF file name is required")
            else:
                with tempfile.TemporaryDirectory() as tmpdir:
                    text = docx_to_text(ACTIVE_TEMPLATE)
                    final_text = replace_placeholders(text, user_data)

                    pdf_path = os.path.join(tmpdir, f"{pdf_name}.pdf")
                    generate_pdf(final_text, pdf_path)

                    with open(pdf_path, "rb") as f:
                        st.download_button(
                            "⬇️ Download PDF",
                            f,
                            file_name=f"{pdf_name}.pdf",
                            mime="application/pdf"
                        )

                st.success("✅ PDF generated successfully")
